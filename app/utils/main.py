from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn

from app.models.data import Data


def exe_docx(id):
    info = {}
    document = Document()
    info = Data.query.filter_by(id=id).first()
    a = document.add_paragraph('河南省高等学校计算机教育研究会职业教育专委会')
    b = document.add_paragraph('常务委员推荐表')
    # 调节第一行字体
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(16)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    a.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 调节第二行字体
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    document.styles['Normal'].font.size = Pt(16)
    document.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
    b.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 循环列表个数
    table = document.add_table(rows=12, cols=8, style='Table Grid')
    for i in range(len(table.rows)):
        table.rows[i].height = Cm(1.5)
        for j in range(len(table.columns)):
            table.cell(i, j).width = Cm(2)
            table.cell(i, j).vertical_alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    # 合并
    table.cell(0, 6).merge(table.cell(3, 7))
    table.cell(1, 1).merge(table.cell(1, 3))
    table.cell(2, 1).merge(table.cell(2, 3))
    table.cell(3, 1).merge(table.cell(3, 2))
    table.cell(3, 4).merge(table.cell(3, 5))
    table.cell(4, 1).merge(table.cell(4, 3))
    table.cell(4, 5).merge(table.cell(4, 7))
    table.cell(5, 0).merge(table.cell(7, 0))
    table.cell(8, 0).merge(table.cell(11, 0))
    table.cell(5, 1).merge(table.cell(7, 7))
    table.cell(8, 1).merge(table.cell(11, 7))
    # 第一行数据
    hdr_cells0 = table.rows[0].cells
    hdr_cells0[0].add_paragraph('姓名')
    hdr_cells0[1].add_paragraph(info.name if info.name else '')
    hdr_cells0[2].add_paragraph('性别')
    hdr_cells0[3].add_paragraph(info.age if info.age else '')
    hdr_cells0[4].add_paragraph('年龄')
    hdr_cells0[5].add_paragraph(info.sex if info.sex else '')
    cell = table.cell(0, 6)
    c_p1 = cell.paragraphs[0]
    c_p1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    c_run1 = c_p1.add_run()
    c_run1.add_picture('..' + str(info.photo) if str(info.photo) else 'app/upload/images/1.jpg', width=Cm(4),
                       height=Cm(6.5))

    # 第二行数据
    hdr_cells1 = table.rows[1].cells
    hdr_cells1[0].add_paragraph('工作单位')
    hdr_cells1[1].add_paragraph(info.work_place if info.work_place else '')
    hdr_cells1[4].add_paragraph('职称/职务')
    hdr_cells1[5].add_paragraph(info.position if info.position else '')
    # 第三行数据
    hdr_cells2 = table.rows[2].cells
    hdr_cells2[0].add_paragraph('毕业学校')
    hdr_cells2[1].add_paragraph(info.school if info.school else '')
    hdr_cells2[4].add_paragraph('专业')
    hdr_cells2[5].add_paragraph(info.major if info.major else '')
    # 第四行
    hdr_cells3 = table.rows[3].cells
    hdr_cells3[0].add_paragraph('学历')
    hdr_cells3[1].add_paragraph(info.education if info.education else '')
    hdr_cells3[3].add_paragraph('学位')
    hdr_cells3[4].add_paragraph(info.degree if info.degree else '')
    # 第五行
    hdr_cells4 = table.rows[4].cells
    hdr_cells4[0].add_paragraph('邮箱/微信号')
    hdr_cells4[1].add_paragraph(info.email if info.email else '')
    hdr_cells4[4].add_paragraph('电话')
    hdr_cells4[5].add_paragraph(info.tel if info.tel else '')
    # 第六行
    hdr_cells5 = table.rows[5].cells
    hdr_cells5[0].add_paragraph('主要工作经历')
    hdr_cells5[1].add_paragraph(info.work_experience if info.work_experience else '')
    # 第8行
    hdr_cells8 = table.rows[8].cells
    hdr_cells8[0].add_paragraph('社会兼职')
    hdr_cells8[1].add_paragraph(info.job if info.job else '')
    # 最后修改文本中所有的字体格式
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(10)
                    font.name = '楷体'
    # 打印列表
    a = '1'
    return document.save(str(id) + str(info.name) + '.docx')
